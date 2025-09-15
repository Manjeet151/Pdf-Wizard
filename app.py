from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
import logging
from docx2pdf import convert as docx2pdf_convert
from PIL import Image
import img2pdf
from fpdf import FPDF
from docx import Document
import tempfile
import sys
import os
import pandas as pd
import time  # Added missing import

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Define allowed extensions
ALLOWED_EXTENSIONS = {
    'word': ['.doc', '.docx'],
    'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
    'excel': ['.xls', '.xlsx', '.csv'],
    'text': ['.txt']
}

def allowed_file(filename, file_type):
    """Check if the file extension is allowed for the given type"""
    if '.' not in filename:
        return False
    ext = filename.lower().rsplit('.', 1)[1]
    return f'.{ext}' in ALLOWED_EXTENSIONS.get(file_type, [])

def convert_doc_to_pdf(input_path, output_path):
    try:
        # Use docx2pdf for DOCX files (works on both Windows and Linux)
        if input_path.lower().endswith('.docx'):
            docx2pdf_convert(input_path, output_path)
            return True
        else:
            # For .doc files, we'll need to use a different approach
            # since docx2pdf doesn't support .doc files directly
            return convert_doc_to_pdf_fallback(input_path, output_path)
    except Exception as e:
        logger.error(f"Word conversion failed: {str(e)}")
        # Fallback to alternative method
        return convert_doc_to_pdf_fallback(input_path, output_path)

def convert_doc_to_pdf_fallback(input_path, output_path):
    """
    Fallback method for document conversion that doesn't rely on Windows COM
    This is a simplified version - you might need to expand this based on your needs
    """
    try:
        # For DOC files, we can try to read with python-docx (might not work for .doc)
        # or use a different approach
        if input_path.lower().endswith('.docx'):
            # Try to use python-docx to read and then create PDF
            document = Document(input_path)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    pdf.multi_cell(0, 10, paragraph.text)
                    pdf.ln(2)
            
            pdf.output(output_path)
            return True
        else:
            # For .doc files, we might need to use a service or different library
            raise Exception("DOC file conversion not supported on this platform. Please convert to DOCX first.")
            
    except Exception as e:
        logger.error(f"Fallback conversion also failed: {str(e)}")
        raise Exception(f"Document conversion failed: {str(e)}")

def convert_word_to_pdf_fallback(input_path, output_path):
    """Alternative fallback for Word conversion"""
    try:
        # Simple text extraction approach
        if input_path.lower().endswith('.docx'):
            document = Document(input_path)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    pdf.multi_cell(0, 10, paragraph.text[:200])  # Limit text length
                    pdf.ln(2)
            
            pdf.output(output_path)
            return True
        else:
            raise Exception("Only DOCX files supported in fallback mode")
    except Exception as e:
        logger.error(f"Word fallback conversion failed: {str(e)}")
        raise Exception(f"Word conversion failed: {str(e)}")

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Max size is 16MB.'}), 413

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    file_type = request.form.get('type')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file_type or file_type not in ALLOWED_EXTENSIONS:
        return jsonify({'error': 'Invalid file type specified'}), 400
    
    if file and allowed_file(file.filename, file_type):
        input_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
        output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        try:
            file.save(input_file.name)
            input_file.close()
            
            # Word conversion
            if file_type == 'word':
                if file.filename.lower().endswith('.docx'):
                    try:
                        docx2pdf_convert(input_file.name, output_file.name)  # Fixed function name
                    except Exception as e:
                        logger.warning(f"docx2pdf failed, using fallback: {str(e)}")
                        convert_word_to_pdf_fallback(input_file.name, output_file.name)
                elif file.filename.lower().endswith('.doc'):
                    convert_doc_to_pdf(input_file.name, output_file.name)
                else:
                    raise Exception("Unsupported Word format")
            
            # Image conversion
            elif file_type == 'image':
                with open(output_file.name, "wb") as f:
                    try:
                        img = Image.open(input_file.name)
                        if img.mode in ('RGBA', 'P'):
                            img = img.convert('RGB')
                        img.save(f, "PDF", resolution=100.0)
                    except Exception as e:
                        logger.warning(f"PIL conversion failed, using img2pdf: {str(e)}")
                        with open(input_file.name, "rb") as img_file:
                            f.write(img2pdf.convert(img_file))
            
            # Excel / CSV conversion
            elif file_type == 'excel':
                try:
                    if input_file.name.endswith('.csv'):
                        df = pd.read_csv(input_file.name)
                    else:
                        df = pd.read_excel(input_file.name)
                    
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=8)
                    
                    col_count = len(df.columns)
                    if col_count > 0:
                        col_width = 190 / min(col_count, 8)  # Adjust for more columns
                        for col in df.columns:
                            pdf.cell(col_width, 10, str(col)[:15], border=1)
                        pdf.ln()
                        for index, row in df.head(50).iterrows():
                            for col in df.columns:
                                pdf.cell(col_width, 10, str(row[col])[:20], border=1)
                            pdf.ln()
                    
                    pdf.output(output_file.name)
                except Exception as e:
                    raise Exception(f"Excel conversion error: {str(e)}")
            
            # Text conversion
            elif file_type == 'text':
                try:
                    with open(input_file.name, 'r', encoding='utf-8', errors='ignore') as txt_file:
                        text_content = txt_file.read()
                    
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    for line in text_content.split('\n')[:100]:
                        if line.strip():
                            pdf.multi_cell(0, 10, line[:500])
                            pdf.ln(2)
                    pdf.output(output_file.name)
                except Exception as e:
                    raise Exception(f"Text conversion error: {str(e)}")
            
            output_file.close()
            original_filename = os.path.splitext(file.filename)[0]
            download_filename = f"{original_filename}.pdf"
            
            return send_file(
                output_file.name,
                as_attachment=True,
                download_name=download_filename,
                mimetype='application/pdf'
            )
        
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            try: 
                os.unlink(input_file.name)
            except: 
                pass
            try: 
                os.unlink(output_file.name)
            except: 
                pass
            return jsonify({'error': str(e)}), 500
        
        finally:
            try: 
                os.unlink(input_file.name)
            except: 
                pass
            try: 
                os.unlink(output_file.name)
            except: 
                pass
    
    return jsonify({'error': 'Invalid file type or format'}), 400

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/health')
def health():
    return jsonify({
        'status': 'healthy', 
        'timestamp': time.time(),
        'service': 'PDF Converter API'
    })

if __name__ == '__main__':
    # Create templates and static directories if they don't exist
    os.makedirs('templates', exist_ok=True)
    os.makedirs('static', exist_ok=True)
    app.run(debug=True, port=5000, host='0.0.0.0')
